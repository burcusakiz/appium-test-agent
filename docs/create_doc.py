#!/usr/bin/env python3
"""
Appium Test Agent - Mimarisi Dokümanı Oluşturucu
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==========================================
    # TITLE PAGE
    # ==========================================
    # Logo ekleme - optional (logo dosyası yoksa atlanır)
    # doc.add_picture('logo.png', width=Inches(3))
    doc.add_paragraph()

    title = doc.add_paragraph('Appium Test Agent')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph('Teknik Mimari Dokümantasyonu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.italic = True

    doc.add_paragraph()

    info = doc.add_paragraph('Tasarım ve Geliştirme | Test Otomasyon Çerçevesi')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in info.runs:
        run.font.size = Pt(12)

    date = doc.add_paragraph('Mayıs 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date.runs:
        run.font.size = Pt(11)
        run.font.italic = True

    # Table of Contents
    doc.add_page_break()
    toc_title = doc.add_paragraph('İçindekiler')
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    doc.add_paragraph()

    toc_entries = [
        ('1. Proje Özeti', '1'),
        ('2. Sistem Mimarisi', '2'),
        ('3. Bileşen Detayları', '3'),
        ('4. Test Akışı', '4'),
        ('5. Teknoloji Stack', '5'),
        ('6. Dizin Yapısı', '6'),
        ('7. Konfigürasyon', '7'),
        ('8. Kullanım Kılavuzu', '8'),
        ('9. Örnek Test', '9'),
        ('10. Raporlama', '10'),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        run1 = p.add_run(entry)
        run1.font.size = Pt(11)
        leader = p.add_run('\t\t\t\t\t\t\t\t\t\t')
        leader.font.size = Pt(11)
        run2 = p.add_run(page)
        run2.font.size = Pt(11)

    # ==========================================
    # PAGE 1: PROJE ÖZETI
    # ==========================================
    doc.add_page_break()
    heading1 = doc.add_paragraph('1. Proje Özeti')
    for run in heading1.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    p1 = doc.add_paragraph(
        'Appium Test Agent, mobil uygulamaların (özellikle Android) ve REST API '
        'servislerinin otomatik testini gerçekleştiren kapsamlı bir test frameworküdür. '
        'Proje, Behavior-Driven Development (BDD) prensiplerine göre tasarlanmıştır.'
    )

    doc.add_paragraph()

    goals = doc.add_paragraph('Amaçlar:')
    for run in goals.runs:
        run.font.size = Pt(11)
        run.font.bold = True

    goals_list = [
        'Mobil uygulama UI testlerini otomatize etmek',
        'REST API testlerini gerçekleştirmek',
        'Test sonuçlarını görselleştirmek ve raporlamak',
        'Test senaryolarını insan okunabilir şekilde dokümante etmek'
    ]

    for goal in goals_list:
        p = doc.add_paragraph(goal)
        p.style = 'List Bullet'

    # ==========================================
    # PAGE 2: SİSTEM MİMARİSİ
    # ==========================================
    doc.add_page_break()
    heading2 = doc.add_paragraph('2. Sistem Mimarisi')
    for run in heading2.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    arch_text = (
        'Aşağıdaki diyagram, test frameworkünün genel mimarisini ve '
        'bileşenler arası bağımlılıkları göstermektedir.'
    )
    doc.add_paragraph(arch_text)

    doc.add_paragraph()

    arch_box = """
+---------------------------------------------------------------+
|                        TEST LAYER                             |
+---------------------------------------------------------------+
|                                                               |
|  +-------------------+       +-------------------+            |
|  |   Appium Tests    |       |   API Tests       |            |
|  |  (UI - Android)   |       |  (REST - Flask)   |            |
|  +--------+----------+       +-------------------+            |
|           |                                                   |
|  +--------v----------+                                        |
|  |   Page Objects    |                                        |
|  |  (POM Pattern)    |                                        |
|  +--------+----------+                                        |
|           |                                                   |
|  +--------v----------+       +-------------------+            |
|  |   Appium Driver   |       |   Requests        |            |
|  +--------+----------+       +-------------------+            |
|           |                                                   |
|  +--------v--------------------------------------+            |
|  |              Test Runner (pytest)             |            |
|  +---------------------+-------------------------+            |
|                       |                                       |
|             +---------v---------+                             |
|             |   HTML Reports    |                             |
|             +-------------------+                             |
+---------------------------------------------------------------+
"""
    p = doc.add_paragraph(arch_box)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph(
        'Mimari, üç ana katmandan oluşur:'
    )

    layers = [
        ('Test Katmanı', 'Appium Tests ve API Tests bileşenleri'),
        ('Arka Plan Katmanı', 'Page Objects ve Driver yönetimi'),
        ('Test Çalıştırma Katmanı', 'pytest framework ve raporlama')
    ]

    for layer, desc in layers:
        p = doc.add_paragraph(f'{layer} - {desc}')
        p.style = 'List Bullet'

    # ==========================================
    # PAGE 3: BILEŞEN DETAYLARI - MOCK API
    # ==========================================
    doc.add_page_break()
    heading3 = doc.add_paragraph('3. Bileşen Detayları')
    for run in heading3.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    heading3_1 = doc.add_paragraph('3.1 Mock API Server')
    for run in heading3_1.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_paragraph()

    mock_desc = (
        'Flask tabanlı bir REST API simülasyon sunucusudur. Mobil '
        'uygulama testleri sırasında gerçekte backend servisine '
        'ihtiyac duymadan test ortamı oluşturmayı sağlar.'
    )
    doc.add_paragraph(mock_desc)

    doc.add_paragraph()

    doc.add_paragraph('Özellikler:')
    features = [
        ('Port', '8080'),
        ('Endpoints', '/api/login, /api/health'),
        ('Response Format', 'JSON'),
        ('Authentication', 'Token-based'),
    ]

    for feat, val in features:
        p = doc.add_paragraph(f'{feat}: {val}')
        p.style = 'List Bullet'

    # ==========================================
    # PAGE 4: BILEŞEN DETAYLARI - APPIUM
    # ==========================================
    doc.add_page_break()

    heading3_2 = doc.add_paragraph('3.2 Appium Driver')
    for run in heading3_2.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_paragraph()

    appium_desc = (
        'Android uygulamalarının UI testlerini gerçekleştirmek için '
        'Appium frameworkunu kullanır. Appium Servera RPC çağrıları '
        'ile bağlanır.'
    )
    doc.add_paragraph(appium_desc)

    doc.add_paragraph()

    doc.add_paragraph('Özellikler:')
    features = [
        ('Port', '4723'),
        ('Uyumluluk', 'Appium Server 2.x'),
        ('Otomatik Driver', 'conftest.py ile yönetilir'),
        ('Implicit Wait', '10 saniye'),
        ('No Reset', 'False (temiz başlangıç)'),
    ]

    for feat, val in features:
        p = doc.add_paragraph(f'{feat}: {val}')
        p.style = 'List Bullet'

    doc.add_paragraph()

    heading3_3 = doc.add_paragraph('3.3 Page Object Model (POM)')
    for run in heading3_3.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_paragraph()

    pom_desc = (
        'Kod tekrarini azaltmak ve test kodunun okunabilirliğini '
        'artırmak için Page Object Pattern kullanılmaktadır. Her '
        'ekran için bir class tanımlanır ve locatorlar tek yerde '
        'toplanır.'
    )
    doc.add_paragraph(pom_desc)

    code_snippet = '''
class LoginPage:
    # Locators
    USERNAME_FIELD = (AppiumBy.ID, "com.demo.loginapp:id/username_input")
    PASSWORD_FIELD = (AppiumBy.ID, "com.demo.loginapp:id/password_input")
    LOGIN_BUTTON = (AppiumBy.ID, "com.demo.loginapp:id/login_button")

    def __init__(self, driver):
        self.driver = driver

    def login(self, username, password):
        self.enter_username(username)
        self.enter_password(password)
        self.click_login()
'''
    p = doc.add_paragraph(code_snippet)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # ==========================================
    # PAGE 5: BILEŞEN DETAYLARI - PYTEST
    # ==========================================
    doc.add_page_break()

    heading3_4 = doc.add_paragraph('3.4 Test Runner (pytest)')
    for run in heading3_4.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_paragraph()

    pytest_desc = (
        'Test discovery ve çalıştırması için pytest kullanılmaktadır. '
        'run_tests.py scripti, pytest komutlarını parametrelerle beraber '
        'çalıştırır ve HTML rapor üretir.'
    )
    doc.add_paragraph(pytest_desc)

    doc.add_paragraph()

    doc.add_paragraph('Özellikler:')
    features = [
        ('Test Discovery', 'tests/ dizini altındaki test_*.py dosyaları'),
        ('Raporlama', 'pytest-html ile HTML rapor'),
        ('Otomatik Screenshot', 'Hata durumunda captured'),
        ('Fixture Yönetimi', 'conftest.py ile ortak fixturelar'),
        ('Parallel Execution', 'pytest-xdist ile paralel test'),
    ]

    for feat, val in features:
        p = doc.add_paragraph(f'{feat}: {val}')
        p.style = 'List Bullet'

    # ==========================================
    # PAGE 6: TEST AKIŞI
    # ==========================================
    doc.add_page_break()
    heading4 = doc.add_paragraph('4. Test Akışı')
    for run in heading4.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        'Testlerin çalışması için gerekli adımlar aşağıdaki gibidir:'
    )

    flow_steps = [
        '1. Mock Server Başlatılır (python3 mock_api_server.py)',
        '2. Appium Server Başlatılır (appium --port 4723)',
        '3. Test Senaryosu Seçilir (test-scenarios/)',
        '4. Test Kodu Çalıştırılır (pytest / make test)',
        '5. Sonuç Raporlanır (reports/YYYY-MM-DD/)',
        '6. Screenshot ve Log Kaydedilir',
    ]

    for step in flow_steps:
        p = doc.add_paragraph(step)
        p.style = 'List Number'

    doc.add_paragraph()

    doc.add_paragraph(
        'Aşağıdaki akış diyagramı, test senaryosunun nasıl oluşturulduğunu '
        've çalıştırıldığını göstermektedir:'
    )

    flow_chart = """
+---------+     +---------+     +------------+
| BDD     |---->| Test    |---->| Automation |
| Senaryo |     | Code    |     | Code       |
| (.md)   |     | (.py)   |     | (.py)      |
+---------+     +---------+     +------------+
      |               |                |
      |               |                |
      +-------+-------+----------------+
            |
      +-------v-------+
      |   Test Run    |
      |   (pytest)    |
      +-------+-------+
            |
      +-------v-------+
      |  HTML Report  |
      |  + Screens    |
      +---------------+
"""
    p = doc.add_paragraph(flow_chart)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    # ==========================================
    # PAGE 7: TEKNOLOJI STACK
    # ==========================================
    doc.add_page_break()
    heading5 = doc.add_paragraph('5. Teknoloji Stack')
    for run in heading5.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(7)
    table.columns[2].width = Cm(5)

    hdr = table.rows[0].cells
    hdr[0].text = 'Kategori'
    hdr[1].text = 'Teknoloji'
    hdr[2].text = 'Versiyon'

    tech_data = [
        ('Testing Framework', 'Pytest', '7.4+'),
        ('Mobile Testing', 'Appium', '4.0+'),
        ('API Testing', 'Requests', '2.31+'),
        ('Web Framework', 'Flask', '2.0+'),
        ('Reporting', 'pytest-html', '4.1+'),
        ('Selenium', 'Selenium', '4.15+'),
        ('Image Processing', 'Pillow', '10.0+'),
        ('Utilities', 'python-dotenv', '1.0+'),
        ('Parallel Run', 'pytest-xdist', '3.3+'),
    ]

    for category, tech, version in tech_data:
        row = table.add_row().cells
        row[0].text = category
        row[1].text = tech
        row[2].text = version

    # ==========================================
    # PAGE 8: DIZIN YAPISI
    # ==========================================
    doc.add_page_break()
    heading6 = doc.add_paragraph('6. Dizin Yapısı')
    for run in heading6.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    tree = """
appium-test-agent/
├── apps/                              ← APK dosyaları
├── dashboard-app/                     ← Dashboard web uygulaması
│   ├── app.py                         ← Flask sunucu
│   ├── templates/                     ← HTML şablonlar
│   └── static/                        ← Statik dosyalar
├── docs/                              ← Dokümantasyon
│   ├── ARCHITECTURE.md
│   ├── WORKFLOW.md
│   └── CONFIGURATION.md
├── mock_api_server.py                 ← Flask mock REST API
├── run_tests.py                       ← Test runner scripti
├── test-scenarios/                    ← BDD senaryoları
│   └── TS-0001-LOGIN/
│       ├── TC-0001-valid-login.md
│       └── ...
├── tests/                             ← Test otomasyon
│   ├── conftest.py
│   ├── android/login/
│   ├── api/
│   └── page_objects/
├── reports/                           ← Raporlar
│   └── YYYY-MM-DD/
├── .claude/                           ← Claude Code agent/command/skill dosyaları
│   ├── agents/
│   ├── commands/
│   └── skills/
├── Makefile                           ← Build automation
├── pytest.ini                         ← Pytest config
├── requirements.txt                   ← Dependencies
└── CLAUDE.md                          ← Project instructions
"""
    p = doc.add_paragraph(tree)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # ==========================================
    # PAGE 9: KONFIGURASYON
    # ==========================================
    doc.add_page_break()
    heading7 = doc.add_paragraph('7. Konfigürasyon')
    for run in heading7.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    heading7_1 = doc.add_paragraph('7.1 Pytest Konfigürasyonu (pytest.ini)')
    for run in heading7_1.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    pytest_config = '''
[pytest]
addopts = -v --tb=short --strict-markers
testpaths = tests
python_files = test_*.py
python_classes = Test*
python_functions = test_*

markers =
    smoke: Quick sanity tests
    regression: Full regression suite
    android: Android UI tests
    api: API tests
    login: Login feature tests
'''
    p = doc.add_paragraph(pytest_config)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # ==========================================
    # PAGE 10: KULLANIM
    # ==========================================
    doc.add_page_break()
    heading8 = doc.add_paragraph('8. Kullanım Kılavuzu')
    for run in heading8.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    doc.add_paragraph('Aşağıdaki komutlar ile test ortamı hızlıca kurulabilir:')

    commands = [
        ('Mock Server Başlat:', 'make mock'),
        ('Appium Server Başlat:', 'make appium'),
        ('Tüm Testleri Çalıştır:', 'make test'),
        ('Android Testleri:', 'make test-android'),
        ('API Testleri:', 'make test-api'),
        ('Dashboard Başlat:', 'make dashboard'),
    ]

    for cmd_desc, cmd in commands:
        p = doc.add_paragraph(f'{cmd_desc} {cmd}')
        p.style = 'List Bullet'

    doc.add_paragraph()

    heading8_1 = doc.add_paragraph('8.1 Dashboard Özellikleri')
    for run in heading8_1.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_paragraph()

    dashboard_features = [
        'Test senaryoları listesi',
        'Seçili testleri çalıştırma',
        'Live emulator ekranı (SSE)',
        'Test sonuç logu',
        'Rapor yönetimi',
    ]

    for feat in dashboard_features:
        p = doc.add_paragraph(feat)
        p.style = 'List Bullet'

    # ==========================================
    # PAGE 11: ÖRNEK TEST
    # ==========================================
    doc.add_page_break()
    heading9 = doc.add_paragraph('9. Örnek Test')
    for run in heading9.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    heading9_1 = doc.add_paragraph('9.1 BDD Senaryo')
    for run in heading9_1.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    bdd_scenario = '''
# TC-0001: Valid Login

App: com.demo.loginapp
Suite: TS-0001-LOGIN
Priority: High

Steps:
1. Launch the application
2. Enter username: testuser
3. Enter password: password123
4. Tap the Login button

Expected:
Home screen is displayed. Welcome message contains "testuser".

Automation: tests/android/login/test_TC001_valid_login.py
'''
    p = doc.add_paragraph(bdd_scenario)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    heading9_2 = doc.add_paragraph('9.2 Otomasyon Kodu')
    for run in heading9_2.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    auto_code = '''
import pytest

class TestTC001:
    @pytest.mark.smoke
    @pytest.mark.login
    def test_valid_login(self, driver):
        from tests.page_objects.login_page import LoginPage

        login_page = LoginPage(driver)
        login_page.login("testuser", "password123")

        assert login_page.is_logged_in()
'''
    p = doc.add_paragraph(auto_code)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # ==========================================
    # PAGE 12: RAPORLAMA
    # ==========================================
    doc.add_page_break()
    heading10 = doc.add_paragraph('10. Raporlama')
    for run in heading10.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        'Her test çalıştırmasından sonra otomatik olarak raporlar '
        'reports/YYYY-MM-DD/ dizinine kaydedilir:'
    )

    report_features = [
        'report.html - HTML test raporu',
        'screenshots/ - Hata anı görüntüleri',
        'Self-contained HTML - Dışa aktarım için hazır',
    ]

    for feat in report_features:
        p = doc.add_paragraph(feat)
        p.style = 'List Bullet'

    doc.add_paragraph()

    doc.add_paragraph('Screenshot Naming Convention:')
    p_code = doc.add_paragraph('Format: screenshot_{test_name}_{YYYYMMDD_HHMMSS}.png')
    for run in p_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    doc.add_paragraph()

    screenshot_example = 'Örnek: screenshot_test_valid_login_20260502_143025.png'
    p = doc.add_paragraph(screenshot_example)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = 'Appium Test Agent - Mimarisi v1.0 | Mayıs 2026'
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(9)

    # Save
    output_path = '/Users/semih.sakiz/DEVELOPER/appium-test-agent/docs/Appium_Test_Agent_Mimarisi.docx'
    doc.save(output_path)

    size_bytes = os.path.getsize(output_path)
    print(f'Doküman oluşturuldu: {output_path}')
    print(f'Boyut: {size_bytes} bytes ({size_bytes / 1024:.2f} KB)')

if __name__ == '__main__':
    create_document()
